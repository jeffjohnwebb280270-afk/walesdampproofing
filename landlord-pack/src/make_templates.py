#!/usr/bin/env python3
"""Emit each template from templates.py as both .docx and print-ready HTML."""
import html
import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Cm, RGBColor

import templates as T

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, '..', 'build')
INK, MID, RED = RGBColor(0x15, 0x13, 0x1B), RGBColor(0x5F, 0x5A, 0x68), RGBColor(0xC8, 0x10, 0x2E)
MARK = ('<svg viewBox="0 0 60 76" aria-hidden="true"><path d="M8 6h16l6 10 6-10h16v52a12 '
        '12 0 0 1-12 12H20A12 12 0 0 1 8 58V6z"/></svg>')


# ---------------------------------------------------------------- docx
def _shade(cell, fill):
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear')
    el.set(qn('w:fill'), fill)
    cell._tc.get_or_add_tcPr().append(el)


def _run(p, text, size=10, bold=False, italic=False, colour=None, caps=False):
    r = p.add_run(text.upper() if caps else text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if colour is not None:
        r.font.color.rgb = colour
    return r


def as_docx(spec):
    d = Document()
    s = d.sections[0]
    s.top_margin = s.bottom_margin = Cm(1.7)
    s.left_margin = s.right_margin = Cm(1.6)
    n = d.styles['Normal']
    n.font.name, n.font.size, n.font.color.rgb = 'Calibri', Pt(10), INK
    n.paragraph_format.space_after = Pt(5)

    p = d.add_paragraph(); _run(p, 'WALES DAMP PROOFING', 9, bold=True, colour=MID)
    p.paragraph_format.space_after = Pt(1)
    p = d.add_paragraph(); _run(p, spec['title'], 17, bold=True)
    p.paragraph_format.space_after = Pt(3)
    p = d.add_paragraph(); _run(p, spec['strap'], 9.5, colour=MID)
    p.paragraph_format.space_after = Pt(11)

    for headline, kind, payload in spec['sections']:
        if headline:
            p = d.add_paragraph()
            _run(p, headline, 8.5, bold=True, colour=RED, caps=True)
            p.paragraph_format.space_before = Pt(11)
            p.paragraph_format.space_after = Pt(3)
        if kind == 'fields':
            t = d.add_table(rows=len(payload), cols=2)
            t.style = 'Table Grid'
            for i, (label, hint) in enumerate(payload):
                c0, c1 = t.rows[i].cells
                c0.width, c1.width = Cm(5.4), Cm(12.4)
                _shade(c0, 'F4F3F6')
                c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                _run(c0.paragraphs[0], label, 9, bold=True)
                if hint:
                    _run(c1.paragraphs[0], hint, 8, italic=True, colour=MID)
                c1.paragraphs[0].paragraph_format.space_after = Pt(7)
        elif kind == 'checks':
            cols = 3
            t = d.add_table(rows=(len(payload) + cols - 1) // cols, cols=cols)
            for i, item in enumerate(payload):
                _run(t.rows[i // cols].cells[i % cols].paragraphs[0], '☐  ' + item, 9)
        elif kind == 'note':
            p = d.add_paragraph(); _run(p, payload, 7.8, colour=MID)
            p.paragraph_format.space_before = Pt(10)
        elif kind == 'letter':
            for text, style in payload:
                p = d.add_paragraph()
                if text:
                    _run(p, text, 9.5 if style == 'grey' else 10,
                         bold=(style == 'bold'),
                         colour=MID if style == 'grey' else None)
                p.paragraph_format.space_after = Pt(8)

    p = d.add_paragraph(); _run(p, T.DISCLAIMER, 7.4, colour=MID)
    p.paragraph_format.space_before = Pt(14)

    fp = d.sections[0].footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(fp, f'Wales Damp Proofing  ·  Jeff Webb PCAQT 957  ·  07446 522034  '
             f'·  walesdampproofing.co.uk  ·  {spec["stem"][:2]} {spec["title"]}',
         7, colour=MID)

    d.save(os.path.join(OUT, spec['stem'] + '.docx'))


# ---------------------------------------------------------------- html
def as_html(spec):
    e = html.escape
    out = [f'''<!doctype html><html lang="en-GB"><head><meta charset="utf-8">
<title>{e(spec["title"])}</title><link rel="stylesheet" href="../src/pack.css">
<style>
  .fieldrow{{display:flex;border:1px solid var(--line);border-top:0}}
  .fieldrow:first-of-type{{border-top:1px solid var(--line)}}
  .fieldrow .l{{width:46mm;flex:none;background:var(--wash);
    border-right:1px solid var(--line);padding:7px 9px;
    font-family:var(--f-display);font-weight:700;font-size:8.6pt}}
  .fieldrow .v{{flex:1;padding:7px 9px;min-height:34px}}
  .fieldrow .v i{{font-size:7.8pt;color:var(--mid);font-style:italic}}
  .checks{{display:flex;flex-wrap:wrap;margin:8px 0 2px}}
  .checks span{{width:33.33%;font-size:9pt;padding:3px 0}}
  .sec{{font-family:var(--f-mono);font-size:7.6pt;letter-spacing:.14em;
    text-transform:uppercase;color:var(--red);font-weight:700;
    margin:15px 0 5px}}
  .letter p{{margin:0 0 9px}}
  .letter .grey{{color:var(--mid)}}
</style></head><body>
<div class="mast">{MARK}
  <div class="who"><b>Wales Damp Proofing</b><span>Damp &amp; timber &middot; PCA qualified</span></div>
  <div class="contact">07446 522034<br>contact@walesdampproofing.co.uk<br>walesdampproofing.co.uk</div>
</div>
<h1 style="font-size:17pt">{e(spec["title"])}</h1>
<p class="small" style="font-size:9.4pt;margin-bottom:4px">{e(spec["strap"])}</p>''']

    for headline, kind, payload in spec['sections']:
        if headline:
            out.append(f'<div class="sec">{e(headline)}</div>')
        if kind == 'fields':
            for label, hint in payload:
                hint_html = f'<i>{e(hint)}</i>' if hint else ''
                out.append(f'<div class="fieldrow"><div class="l">{e(label)}</div>'
                           f'<div class="v">{hint_html}</div></div>')
        elif kind == 'checks':
            out.append('<div class="checks">' + ''.join(
                f'<span>&#9744;&nbsp; {e(i)}</span>' for i in payload) + '</div>')
        elif kind == 'note':
            out.append(f'<p class="small" style="font-size:7.9pt">{e(payload)}</p>')
        elif kind == 'letter':
            out.append('<div class="letter">')
            for text, style in payload:
                if not text:
                    out.append('<p>&nbsp;</p>')
                elif style == 'bold':
                    out.append(f'<p><strong>{e(text)}</strong></p>')
                else:
                    cls = ' class="grey"' if style == 'grey' else ''
                    out.append(f'<p{cls}>{e(text)}</p>')
            out.append('</div>')

    out.append(f'<div class="note">{e(T.DISCLAIMER)}</div>')
    out.append(f'<div class="foot"><span>Wales Damp Proofing &middot; Jeff Webb PCAQT 957'
               f'</span><span class="r">{spec["stem"][:2]} &middot; {e(spec["title"])}'
               f'</span></div></body></html>')
    open(os.path.join(OUT, spec['stem'] + '.html'), 'w', encoding='utf-8').write('\n'.join(out))


if __name__ == '__main__':
    os.makedirs(OUT, exist_ok=True)
    for spec in T.ALL:
        as_docx(spec)
        as_html(spec)
    print(f'{len(T.ALL)} templates -> .docx + .html')
